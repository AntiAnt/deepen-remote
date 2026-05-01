import argparse
import os
import re
import tempfile

import docx
from docx.shared import Pt
from reliquery import Relic
from docx import Document

from pipeline.pipeline import get_pipeline_service


def generate_doc_from_md_string(text: str) -> Document:
    """
    Converts markdown to DOCX with:
    - Proper headings
    - Bold text
    - Bullet & numbered lists
    - Real clickable hyperlinks
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # === Headings ===
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)

        # === Bold ===
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True

        elif "**" in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # === Bullet points ===
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")

        # === Numbered lists ===
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(line, style="List Number")

        # === Links: [Text](URL) ===
        elif "[" in line and "](" in line and line.endswith(")"):
            # Extract link parts
            match = re.search(r"\[(.+?)\]\((.+?)\)", line)
            if match:
                text = match.group(1)
                url = match.group(2)

                p = doc.add_paragraph()
                run = p.add_run(text)
                # Add real clickable hyperlink
                hyperlink = p.add_hyperlink(url, run)
                # Optional: make link blue and underlined
                hyperlink.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # Blue
                hyperlink.font.underline = True
            else:
                doc.add_paragraph(line)

        # === Regular paragraph ===
        else:
            doc.add_paragraph(line)

    return doc


def main():
    parser = argparse.ArgumentParser(
        description="Deepen pipeline: ingest, transcribe, and summarize informational videos.",
        formatter_class=argparse.RawTextHelpFormatter,
    )

    parser.add_argument("--relic-name", required=True)
    parser.add_argument("--relic-type", required=True)
    parser.add_argument("--reliquery-config-root", required=True)

    args = parser.parse_args()
    relic_name = args.relic_name
    relic_type = args.relic_type
    reliquery_config_root = args.reliquery_config_root

    relic = Relic(
        name=relic_name,
        relic_type=relic_type,
        storage_name="remote_s3",
        reliquery_config_root=reliquery_config_root,
    )

    video_metadata = relic.get_json(name="metadata")
    remote_config = relic.get_json(name="remote-config")

    pipeline = get_pipeline_service(remote_config, video_metadata)
    # transcription

    if "whisper-transcript" in relic.list_json():
        transcription = relic.get_json(name="whisper-transcript")
    else:
        transcription = pipeline.transcribe(audio_obj=relic.get_audio(name="audio.wav"))
        relic.add_json(name="whisper-transcript", json_data=transcription)

    if "readable-whisper-transcript" in relic.list_text():
        readable_transcript = relic.get_text(name="readable-whisper-transcript")
    else:
        readable_transcript = pipeline.get_readable_transcript(transcription)
        relic.add_text(name="readable-whisper-transcript", text=readable_transcript)

    # Summary
    summary = pipeline.summarize(transcript=readable_transcript)

    relic.add_text(name=relic_name, text=summary)

    document: Document = generate_doc_from_md_string(text=summary)

    with tempfile.TemporaryDirectory() as temp_dir:
        docx_filename = f"{relic_name}.docx"
        temp_path = os.path.join(temp_dir, docx_filename)

        document.save(temp_path)
        relic.add_files_from_path(name=docx_filename, path=temp_path)

    print("Main: audio processing complete.")


if __name__ == "__main__":
    main()
